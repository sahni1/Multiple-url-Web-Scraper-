import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document

def get_page_urls(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        a_tags = soup.find_all('a')
        urls = []
        for tag in a_tags:
            href = tag.get('href')
            if href:
                full_url = urljoin(url, href)
                urls.append(full_url)
        return urls
    except requests.RequestException as e:
        print(f"Error accessing {url}: {e}")
        return []

def scrape_data(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string if soup.title else 'No title'
        paragraphs = soup.find_all('p')
        content = '\n'.join([p.get_text() for p in paragraphs[:5]])
        return title, content
    except requests.RequestException as e:
        return None, f"Error accessing {url}: {e}"

def save_to_document(results, filename='results.docx'):
    doc = Document()
    doc.add_heading('Scraped Data', 0)
    
    for result in results:
        url, title, content = result
        doc.add_heading(url, level=1)
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    
    doc.save(filename)

if __name__ == "__main__":
    webpage_url = input("Enter the webpage URL: ")
    urls = get_page_urls(webpage_url)
    results = []
    
    for url in urls:
        title, content = scrape_data(url)
        if title is not None:
            results.append((url, title, content))
            print(f"Scraped {url}")
        else:
            print(f"Failed to scrape {url}")
    
    save_to_document(results, filename='file location')

    print("Data saved to results.docx")
